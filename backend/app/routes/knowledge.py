import os
import uuid
from flask import Blueprint, request, current_app
from flask_jwt_extended import jwt_required, get_jwt_identity
from app import db
from app.models.knowledge import KnowledgeDocument, KnowledgeChunk
from app.services.rag_service import split_text, add_to_vector_store, remove_doc_from_store
from app.utils.response import success, error, paginate_response

knowledge_bp = Blueprint('knowledge', __name__)


@knowledge_bp.route('', methods=['GET'])
@jwt_required()
def get_documents():
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 10, type=int)
    keyword = request.args.get('keyword', '').strip()
    doc_type = request.args.get('doc_type', '').strip()
    category = request.args.get('category', '').strip()

    query = KnowledgeDocument.query
    if keyword:
        query = query.filter(
            db.or_(
                KnowledgeDocument.title.like(f'%{keyword}%'),
                KnowledgeDocument.description.like(f'%{keyword}%'),
            )
        )
    if doc_type:
        query = query.filter_by(doc_type=doc_type)
    if category:
        query = query.filter_by(category=category)

    query = query.order_by(KnowledgeDocument.created_at.desc())
    data = paginate_response(query, page, per_page)
    return success(data)


@knowledge_bp.route('/<int:doc_id>', methods=['GET'])
@jwt_required()
def get_document(doc_id):
    doc = KnowledgeDocument.query.get(doc_id)
    if not doc:
        return error('文档不存在', 404)

    result = doc.to_dict()
    result['content'] = doc.content or ''
    result['chunks'] = [c.to_dict() for c in doc.chunks.order_by(KnowledgeChunk.chunk_index).all()]
    return success(result)


@knowledge_bp.route('', methods=['POST'])
@jwt_required()
def create_document():
    user_id = int(get_jwt_identity())
    data = request.get_json()

    if not data.get('title'):
        return error('文档标题不能为空')
    if not data.get('content'):
        return error('文档内容不能为空')

    doc = KnowledgeDocument(
        title=data['title'],
        doc_type=data.get('doc_type', 'guideline'),
        category=data.get('category'),
        content=data['content'],
        description=data.get('description'),
        uploaded_by=user_id,
    )
    db.session.add(doc)
    db.session.commit()

    return success(doc.to_dict(), '知识文档创建成功')


@knowledge_bp.route('/upload', methods=['POST'])
@jwt_required()
def upload_document():
    user_id = int(get_jwt_identity())

    if 'file' not in request.files:
        return error('请选择文件')

    file = request.files['file']
    if not file.filename:
        return error('文件名为空')

    allowed_ext = {'txt', 'pdf', 'docx', 'doc', 'md'}
    ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
    if ext not in allowed_ext:
        return error(f'不支持的文件格式，仅支持: {", ".join(allowed_ext)}')

    upload_dir = current_app.config.get('UPLOAD_FOLDER', './uploads')
    os.makedirs(upload_dir, exist_ok=True)
    filename = f"{uuid.uuid4().hex}.{ext}"
    filepath = os.path.join(upload_dir, filename)
    file.save(filepath)

    content = ''
    try:
        if ext == 'txt' or ext == 'md':
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
        elif ext == 'docx':
            from docx import Document as DocxDocument
            doc_file = DocxDocument(filepath)
            content = '\n'.join([p.text for p in doc_file.paragraphs if p.text.strip()])
        elif ext == 'pdf':
            from PyPDF2 import PdfReader
            reader = PdfReader(filepath)
            content = '\n'.join([page.extract_text() or '' for page in reader.pages])
    except Exception as e:
        return error(f'文件解析失败: {str(e)}')

    title = request.form.get('title', file.filename)
    doc_type = request.form.get('doc_type', 'guideline')
    category = request.form.get('category', '')

    doc = KnowledgeDocument(
        title=title,
        doc_type=doc_type,
        category=category,
        file_path=filepath,
        file_size=os.path.getsize(filepath),
        content=content,
        uploaded_by=user_id,
    )
    db.session.add(doc)
    db.session.commit()

    return success(doc.to_dict(), '文件上传成功')


@knowledge_bp.route('/<int:doc_id>/vectorize', methods=['POST'])
@jwt_required()
def vectorize_document(doc_id):
    """将文档向量化并存入向量库"""
    doc = KnowledgeDocument.query.get(doc_id)
    if not doc:
        return error('文档不存在', 404)

    if not doc.content:
        return error('文档内容为空，无法向量化')

    chunks = split_text(doc.content)
    if not chunks:
        return error('文档切分失败')

    KnowledgeChunk.query.filter_by(document_id=doc_id).delete()

    chunk_records = []
    for i, chunk_text in enumerate(chunks):
        chunk = KnowledgeChunk(
            document_id=doc_id,
            chunk_index=i,
            content=chunk_text,
            token_count=len(chunk_text),
        )
        chunk_records.append(chunk)
        db.session.add(chunk)

    added_count = add_to_vector_store(chunks, doc_id, doc.title)

    doc.is_vectorized = True
    doc.chunk_count = len(chunks)
    db.session.commit()

    return success({
        'chunk_count': len(chunks),
        'vectorized_count': added_count,
    }, f'文档向量化完成，共 {len(chunks)} 个切片，{added_count} 个成功向量化')


@knowledge_bp.route('/<int:doc_id>', methods=['PUT'])
@jwt_required()
def update_document(doc_id):
    doc = KnowledgeDocument.query.get(doc_id)
    if not doc:
        return error('文档不存在', 404)

    data = request.get_json()
    for field in ['title', 'doc_type', 'category', 'content', 'description']:
        if field in data:
            setattr(doc, field, data[field])

    if 'content' in data:
        doc.is_vectorized = False

    db.session.commit()
    return success(doc.to_dict(), '更新成功')


@knowledge_bp.route('/<int:doc_id>', methods=['DELETE'])
@jwt_required()
def delete_document(doc_id):
    doc = KnowledgeDocument.query.get(doc_id)
    if not doc:
        return error('文档不存在', 404)

    if doc.is_vectorized:
        remove_doc_from_store(doc_id)

    if doc.file_path and os.path.exists(doc.file_path):
        os.remove(doc.file_path)

    db.session.delete(doc)
    db.session.commit()
    return success(message='文档删除成功')


@knowledge_bp.route('/categories', methods=['GET'])
@jwt_required()
def get_categories():
    categories = db.session.query(KnowledgeDocument.category).filter(
        KnowledgeDocument.category.isnot(None),
        KnowledgeDocument.category != '',
    ).distinct().all()
    return success([c[0] for c in categories])
