import io
import re
from flask import Blueprint, request, send_file
from flask_jwt_extended import jwt_required, get_jwt_identity
from app import db
from app.models.plan import MedicalPlan
from app.models.patient import MedicalRecord
from app.services.llm_service import generate_medical_plan, analyze_symptoms
from app.services.rag_service import search_similar
from app.utils.response import success, error, paginate_response
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

plan_bp = Blueprint('plan', __name__)


@plan_bp.route('/generate', methods=['POST'])
@jwt_required()
def generate_plan():
    """基于病例信息生成医疗方案（核心功能）"""
    user_id = int(get_jwt_identity())
    data = request.get_json()

    input_text = data.get('input_text', '').strip()
    record_id = data.get('record_id')
    patient_id = data.get('patient_id')
    title = data.get('title', '智慧医疗方案')

    if not input_text and record_id:
        record = MedicalRecord.query.get(record_id)
        if record:
            input_text = record.to_text()
            patient_id = patient_id or record.patient_id

    if not input_text:
        return error('请提供病例信息')

    # RAG检索相关文档
    retrieved_docs, _ = search_similar(input_text, top_k=5)
    context_text = ""
    reference_sources = []
    if retrieved_docs:
        context_parts = []
        for i, doc in enumerate(retrieved_docs):
            context_parts.append(f"【参考{i + 1}】(来源: {doc['doc_title']}, 相关度: {doc['score']:.2f})\n{doc['content']}")
            reference_sources.append(f"{doc['doc_title']}(相关度: {doc['score']:.2f})")
        context_text = '\n\n'.join(context_parts)

    # 调用千问生成方案
    result = generate_medical_plan(input_text, context_text)

    if not result['success']:
        return error(f"方案生成失败: {result['error']}")

    generated_content = result['content']

    plan = MedicalPlan(
        record_id=record_id,
        patient_id=patient_id,
        title=title,
        input_text=input_text,
        retrieved_context=context_text if context_text else None,
        generated_plan=generated_content,
        reference_sources='\n'.join(reference_sources) if reference_sources else None,
        confidence_score=retrieved_docs[0]['score'] if retrieved_docs else None,
        generated_by=user_id,
    )
    db.session.add(plan)
    db.session.commit()

    return success(plan.to_dict(), '医疗方案生成成功')


@plan_bp.route('/analyze', methods=['POST'])
@jwt_required()
def analyze():
    """症状分析"""
    data = request.get_json()
    symptoms = data.get('symptoms', '').strip()

    if not symptoms:
        return error('请提供症状描述')

    result = analyze_symptoms(symptoms)
    if not result['success']:
        return error(f"分析失败: {result['error']}")

    return success({
        'analysis': result['content'],
        'usage': result.get('usage'),
    }, '症状分析完成')


@plan_bp.route('/search', methods=['POST'])
@jwt_required()
def search_knowledge():
    """语义检索知识库"""
    data = request.get_json()
    query = data.get('query', '').strip()
    top_k = data.get('top_k', 5)

    if not query:
        return error('请提供检索内容')

    results, err_msg = search_similar(query, top_k=top_k)
    if err_msg:
        return error(err_msg)
    return success(results, f'检索到 {len(results)} 条相关结果')


@plan_bp.route('', methods=['GET'])
@jwt_required()
def get_plans():
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 10, type=int)
    keyword = request.args.get('keyword', '').strip()
    status = request.args.get('status', '').strip()

    query = MedicalPlan.query
    if keyword:
        query = query.filter(
            db.or_(
                MedicalPlan.title.like(f'%{keyword}%'),
                MedicalPlan.input_text.like(f'%{keyword}%'),
            )
        )
    if status:
        query = query.filter_by(status=status)

    query = query.order_by(MedicalPlan.created_at.desc())
    data = paginate_response(query, page, per_page)
    return success(data)


@plan_bp.route('/<int:plan_id>', methods=['GET'])
@jwt_required()
def get_plan(plan_id):
    plan = MedicalPlan.query.get(plan_id)
    if not plan:
        return error('方案不存在', 404)
    return success(plan.to_dict())


@plan_bp.route('/<int:plan_id>/feedback', methods=['POST'])
@jwt_required()
def submit_feedback(plan_id):
    plan = MedicalPlan.query.get(plan_id)
    if not plan:
        return error('方案不存在', 404)

    data = request.get_json()
    if 'rating' in data:
        plan.rating = data['rating']
    if 'doctor_feedback' in data:
        plan.doctor_feedback = data['doctor_feedback']
    if 'status' in data:
        plan.status = data['status']

    db.session.commit()
    return success(plan.to_dict(), '反馈提交成功')


@plan_bp.route('/<int:plan_id>/export', methods=['GET'])
@jwt_required()
def export_plan(plan_id):
    """导出方案为Word文档"""
    plan = MedicalPlan.query.get(plan_id)
    if not plan:
        return error('方案不存在', 404)

    doc = DocxDocument()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)

    # 标题
    title = doc.add_heading(plan.title or '智慧医疗方案', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息
    doc.add_heading('基本信息', level=2)
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Light Grid Accent 1'
    cells = info_table.rows[0].cells
    cells[0].text = '生成时间'
    cells[1].text = plan.created_at.strftime('%Y-%m-%d %H:%M:%S') if plan.created_at else '-'
    cells = info_table.rows[1].cells
    cells[0].text = '方案状态'
    status_map = {'generated': '已生成', 'reviewed': '已审阅', 'approved': '已采纳', 'rejected': '已驳回'}
    cells[1].text = status_map.get(plan.status, plan.status or '已生成')
    cells = info_table.rows[2].cells
    cells[0].text = '知识库匹配度'
    cells[1].text = f"{plan.confidence_score * 100:.0f}%" if plan.confidence_score else '-'
    doc.add_paragraph()

    # 病例信息
    if plan.input_text:
        doc.add_heading('病例信息', level=2)
        doc.add_paragraph(plan.input_text)

    # 生成方案内容（解析Markdown为纯文本段落）
    if plan.generated_plan:
        doc.add_heading('医疗方案', level=2)
        lines = plan.generated_plan.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            # 处理Markdown标题
            if line.startswith('### '):
                doc.add_heading(line[4:], level=4)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=3)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=2)
            elif line.startswith('- ') or line.startswith('* '):
                # 列表项
                clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line[2:])
                doc.add_paragraph(clean, style='List Bullet')
            elif re.match(r'^\d+\.\s', line):
                clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                doc.add_paragraph(clean, style='List Number')
            else:
                clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                doc.add_paragraph(clean)

    # 参考来源
    if plan.reference_sources:
        doc.add_heading('参考来源', level=2)
        for src in plan.reference_sources.split('\n'):
            if src.strip():
                doc.add_paragraph(src.strip(), style='List Bullet')

    # 医生反馈
    if plan.doctor_feedback:
        doc.add_heading('医生反馈', level=2)
        if plan.rating:
            doc.add_paragraph(f'评分：{plan.rating}/5')
        doc.add_paragraph(plan.doctor_feedback)

    # 免责声明
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run('免责声明：本方案由AI系统辅助生成，仅供医生参考，最终诊疗决策须由执业医师做出。')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(150, 150, 150)

    # 导出
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"{plan.title or '医疗方案'}.docx"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )


@plan_bp.route('/<int:plan_id>', methods=['DELETE'])
@jwt_required()
def delete_plan(plan_id):
    plan = MedicalPlan.query.get(plan_id)
    if not plan:
        return error('方案不存在', 404)
    db.session.delete(plan)
    db.session.commit()
    return success(message='方案删除成功')
